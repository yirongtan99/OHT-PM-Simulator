from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helper: shade a table cell ─────────────────────────────────────────────
def shade_cell(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

# ── Helper: add a heading ──────────────────────────────────────────────────
def add_heading(text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return h

# ── Helper: add a normal paragraph ────────────────────────────────────────
def add_para(text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

# ── Helper: add a file/line reference label ────────────────────────────────
def add_file_ref(filepath, lines):
    p = doc.add_paragraph()
    run_icon = p.add_run("📄  ")
    run_icon.font.size = Pt(10)
    run_path = p.add_run(filepath)
    run_path.bold = True
    run_path.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
    run_path.font.size = Pt(10)
    run_lines = p.add_run(f"  →  Lines {lines}")
    run_lines.italic = True
    run_lines.font.size = Pt(10)
    run_lines.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

# ── Helper: add a side-by-side table (description | code) ─────────────────
def add_side_by_side(description_text, code_text, file_ref=None, lines_ref=None):
    if file_ref:
        add_file_ref(file_ref, lines_ref)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Cm(9)
    table.columns[1].width = Cm(9)

    # Description cell (left, white bg)
    desc_cell = table.cell(0, 0)
    shade_cell(desc_cell, 'F5F8FF')
    desc_para = desc_cell.paragraphs[0]
    desc_run  = desc_para.add_run(description_text)
    desc_run.font.size = Pt(9.5)

    # Code cell (right, dark bg)
    code_cell = table.cell(0, 1)
    shade_cell(code_cell, '1E1E2E')
    code_para = code_cell.paragraphs[0]
    code_run  = code_para.add_run(code_text)
    code_run.font.name  = 'Courier New'
    code_run.font.size  = Pt(8.5)
    code_run.font.color.rgb = RGBColor(0xCD, 0xEF, 0xFF)

    doc.add_paragraph()  # spacer

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_heading('Project ODIN', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x00, 0x5B, 0x96)

sub = doc.add_paragraph('UI & Backend Technical Breakdown')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(16)
sub.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

info = doc.add_paragraph('Author: Tan Yi Rong  |  GlobalFoundries – AMHS Department')
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.runs[0].font.size = Pt(11)
info.runs[0].italic = True

doc.add_paragraph()
note = doc.add_paragraph('📖  How to use this document: Each section shows a plain-English description on the LEFT and the exact code on the RIGHT. The file path and line numbers are shown above each table so you can jump straight to it in VS Code (Ctrl+G → type line number).')
note.runs[0].font.size = Pt(9.5)
note.runs[0].italic = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1: OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
add_heading('1. System Overview', level=1, color='005B96')
add_para(
    'The Project ODIN dashboard is made up of two main files that work together: '
    'the Backend (app.py) which is the "brain" that talks to the physical sensor and runs a mini web server, '
    'and the Frontend (index.html) which is the "face" — the dashboard you see in the browser.',
    size=10
)

# files table
tbl = doc.add_table(rows=3, cols=2)
tbl.style = 'Table Grid'
headers = ['Part', 'File']
data    = [['Backend (Brain)', 'OHT Wheel Measurement Project/app.py'],
           ['Frontend (Face)', 'OHT Wheel Measurement Project/templates/index.html']]
for i, hdr in enumerate(headers):
    cell = tbl.rows[0].cells[i]
    shade_cell(cell, '005B96')
    run = cell.paragraphs[0].add_run(hdr)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        tbl.rows[r+1].cells[c].paragraphs[0].add_run(val).font.size = Pt(9.5)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2: BACKEND — app.py
# ══════════════════════════════════════════════════════════════════════════════
add_heading('2. Backend — app.py', level=1, color='005B96')

# 2.1 Imports
add_heading('2.1  Importing the Tools', level=2)
add_side_by_side(
    'Before the code can do anything, it loads its required tools (called "libraries"). '
    'Think of these like apps on your phone — you open them before using them.\n\n'
    '• socket — The "telephone line" to the physical OMRON sensor.\n'
    '• threading — Runs two tasks at the same time.\n'
    '• flask — Hosts the dashboard as a local website.\n'
    '• csv — Saves data to a spreadsheet file.\n'
    '• datetime — Stamps each reading with the current time.',
    'import socket\nimport time\nimport csv\nimport threading\nfrom datetime import datetime\nfrom collections import deque\nfrom flask import Flask, render_template, jsonify',
    file_ref='app.py', lines_ref='1 – 7'
)

# 2.2 Config
add_heading('2.2  Configuration Settings', level=2)
add_side_by_side(
    'All the key project settings are stored in one place at the top of the file, '
    'making them easy to update without hunting through the code.\n\n'
    '• SENSOR_IP / SENSOR_PORT — Network address of the OMRON ZP-EIP sensor.\n'
    '• CSV_FILENAME — The spreadsheet file where all data is saved permanently.\n'
    '• WHEEL_BASE_DIAMETER_MM — A healthy new wheel is exactly 125mm. The sensor is calibrated to this point.\n'
    '• WHEEL_DETECT_THRESHOLD_MM — If a reading is below 40mm, the code knows a wheel is passing.',
    'SENSOR_IP = \'192.168.250.1\'\nSENSOR_PORT = 64000\nCSV_FILENAME = \'OHT_Wheel_Predictive_Maintenance_Log.csv\'\n\nWHEEL_BASE_DIAMETER_MM = 125.0\nWHEEL_DETECT_THRESHOLD_MM = 40.0',
    file_ref='app.py', lines_ref='21 – 26'
)

# 2.3 Global State
add_heading('2.3  Live Memory Variables', level=2)
add_side_by_side(
    'These are the shared "memory" variables that hold live data.\n\n'
    '• current_distance — Always holds the very latest sensor reading in mm.\n'
    '• distance_history — A rolling list of the last 100 readings, like a snake: '
    'new readings come in at the head and old ones fall off the tail. '
    'This is what feeds the scrolling live graph on the dashboard.',
    'current_distance = 0.0\n\n# Keep the last 100 readings for the live chart\ndistance_history = deque([0.0]*100, maxlen=100)',
    file_ref='app.py', lines_ref='28 – 33'
)

# 2.4 send_omron_command
add_heading('2.4  Talking to the Sensor', level=2)
add_side_by_side(
    'This function sends a command to the OMRON sensor and decodes the reply.\n\n'
    '1. It sends the text "MS,01,0" — asking for the latest measurement.\n'
    '2. The sensor replies in hexadecimal (e.g. FFFF9BCD).\n'
    '3. The code converts hex → integer → millimetres.\n'
    '4. If the sensor has nothing in range, it returns 999.0 as a placeholder.\n\n'
    'The key formula:\n'
    '  int_val / 100000.0\n'
    '…converts the raw integer to millimetres (sensor outputs in 0.01-micrometre steps).',
    'def send_omron_command(sock):\n    command_str = "MS,01,0\\r\\n"\n    sock.sendall(command_str.encode(\'ascii\'))\n    response = sock.recv(1024).decode(\'ascii\').strip()\n\n    parts = response.split(\',\')\n    if len(parts) >= 3 and parts[0] == \'MS\':\n        hex_value_str = parts[2].strip()\n        int_val = int(hex_value_str, 16)\n\n        # Out of range check\n        if 0x7FFF0000 <= int_val <= 0x7FFFFFFF:\n            return 999.0\n        if int_val >= 0x80000000:\n            int_val -= 0x100000000\n\n        mm_value = int_val / 100000.0\n        return mm_value\n    return 999.0',
    file_ref='app.py', lines_ref='39 – 61'
)

# 2.5 Wheel detection loop
add_heading('2.5  Main Sensor Loop — Detecting Wheels', level=2)
add_side_by_side(
    'This is the heart of the system. It runs forever in the background.\n\n'
    '• Lines 75–79: Opens a network connection to the sensor. If it fails, waits 2 seconds and retries.\n'
    '• Line 86: Calls send_omron_command() 100 times per second (every 10ms).\n'
    '• Lines 90–95: If reading is between -40mm and +40mm → wheel is in view! Track the peak.\n'
    '• Lines 97–117: When the reading jumps back above 40mm → wheel has passed!\n'
    '    - Line 104: wheel_diameter = 125mm + peak_reading (the core formula).\n'
    '    - Lines 108–111: Simulates 4 wheels by adding ±0.02mm random noise to the 1 real reading.\n'
    '    - Lines 113–115: Saves one new row to the CSV file.\n'
    '• Line 119: sleep(0.01) — polls 100 times per second.',
    '# Simplified view of the detection logic:\nwhile True:\n    dist = send_omron_command(sock)\n    current_distance = dist\n    distance_history.append(dist)\n\n    if -40.0 < current_distance < 40.0:\n        # Wheel is in view — track peak\n        if not wheel_in_view:\n            wheel_in_view = True\n            max_dist = current_distance\n        if current_distance > max_dist:\n            max_dist = current_distance\n\n    elif wheel_in_view and abs(current_distance) >= 40.0:\n        # Wheel has passed — calculate & save\n        wheel_diameter = WHEEL_BASE_DIAMETER_MM + max_dist\n        fl = round(wheel_diameter, 3)\n        fr = round(wheel_diameter + random.uniform(-0.015, 0.015), 3)\n        # ... save to CSV\n        wheel_in_view = False\n\n    time.sleep(0.01)',
    file_ref='app.py', lines_ref='63 – 123'
)

# 2.6 Flask API
add_heading('2.6  Web Server & API Endpoints', level=2)
add_side_by_side(
    'Flask creates three web "addresses" (called routes) that the browser can call:\n\n'
    '• "/" (Line 125) — Serves the main index.html dashboard page.\n'
    '• "/api/live" (Line 129) — Returns the last 100 live sensor readings as JSON for the scrolling chart.\n'
    '• "/api/history" (Line 143) — Opens the CSV file and returns up to the last 1,000 wheel records as JSON for the Fleet and Analysis tabs.',
    '@app.route(\'/\')\ndef index():\n    return render_template(\'index.html\')\n\n@app.route(\'/api/live\')\ndef api_live():\n    recent = list(distance_history)\n    return jsonify({\n        \'left\': { \'current_distance\': current_distance,\n                  \'history\': recent },\n        \'right\': { \'current_distance\': current_distance,\n                   \'history\': recent }\n    })\n\n@app.route(\'/api/history\')\ndef api_history():\n    # Reads CSV and returns last 1000 records\n    ...',
    file_ref='app.py', lines_ref='125 – 155'
)

# 2.7 Threading startup
add_heading('2.7  Starting Everything Up', level=2)
add_side_by_side(
    'When app.py is run, two workers start simultaneously using Python threading:\n\n'
    '1. A background THREAD runs sensor_polling_thread() — silently reads sensor data 100x/sec forever.\n'
    '2. The Flask WEB SERVER starts on port 5000, waiting for the browser to connect.\n\n'
    'Think of it like two factory workers starting at the same time: one keeps measuring wheels non-stop, '
    'the other stands at the front desk answering browser requests instantly.',
    'if __name__ == \'__main__\':\n    # Start sensor background thread\n    t = threading.Thread(\n        target=sensor_polling_thread,\n        daemon=True\n    )\n    t.start()\n\n    # Start Flask web server\n    app.run(host=\'0.0.0.0\', port=5000, debug=False)',
    file_ref='app.py', lines_ref='157 – 163'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3: FRONTEND — index.html
# ══════════════════════════════════════════════════════════════════════════════
add_heading('3. Frontend — index.html', level=1, color='005B96')

# 3.1 Chart.js
add_heading('3.1  Loading the Charting Library', level=2)
add_side_by_side(
    'This single line loads Chart.js — a free, popular library that draws all the live graphs. '
    'Without this, all the chart areas on the page would just be blank boxes.',
    '<script src="https://cdn.jsdelivr.net/npm/chart.js"></script>',
    file_ref='index.html', lines_ref='8'
)

# 3.2 Theme
add_heading('3.2  The Dark Theme & Color System', level=2)
add_side_by_side(
    'CSS variables define the entire color palette in one place. Changing a variable here '
    'automatically updates every part of the page that uses it.\n\n'
    '• --bg-color: #0b0d17 — The deep navy-black background.\n'
    '• --accent-primary: #00f0ff — The glowing cyan color used for live data.\n'
    '• --accent-secondary: #ff0055 — The red/pink used for warnings and thresholds.',
    ':root {\n    --bg-color: #0b0d17;\n    --panel-bg: rgba(20, 24, 44, 0.6);\n    --accent-primary: #00f0ff;  /* Cyan */\n    --accent-secondary: #ff0055; /* Red */\n    --text-main: #ffffff;\n    --text-muted: #a0a5c0;\n}',
    file_ref='index.html', lines_ref='10 – 18'
)

# 3.3 Pulse animation
add_heading('3.3  The Pulsing "SYSTEM ONLINE" Dot', level=2)
add_side_by_side(
    'The glowing dot in the header uses a CSS animation called "pulse".\n\n'
    '@keyframes defines a mini-animation that makes the dot repeatedly grow (scale 1.2x) '
    'and shrink (scale 0.95x) every 1.5 seconds, creating the heartbeat effect. '
    'This visually communicates that the system is live and actively running.',
    '.live-dot {\n    animation: pulse 1.5s infinite;\n}\n\n@keyframes pulse {\n    0%  { transform: scale(0.95); opacity: 0.8; }\n    50% { transform: scale(1.2);  opacity: 1;   }\n    100%{ transform: scale(0.95); opacity: 0.8; }\n}',
    file_ref='index.html', lines_ref='62 – 75'
)

# 3.4 Tabs
add_heading('3.4  The Three Tab Buttons', level=2)
add_side_by_side(
    'Each button calls the switchTab() JavaScript function (defined at Line 375). '
    'That function simply shows one <div> panel and hides the others by toggling CSS display properties. '
    'There is no page reload — the content swaps instantly.',
    '<button onclick="switchTab(\'live\')">Live Monitoring</button>\n<button onclick="switchTab(\'fleet\')">Fleet Overview</button>\n<button onclick="switchTab(\'oht\')">OHT Analysis</button>\n\n// switchTab logic (Line 375):\nfunction switchTab(tab) {\n    document.getElementById(\'tab-live\').style.display\n        = tab === \'live\' ? \'block\' : \'none\';\n    // ... same for fleet and oht\n}',
    file_ref='index.html', lines_ref='257 – 261 (HTML) & 375 – 386 (JS)'
)

# 3.5 Live Charts
add_heading('3.5  Live Scrolling Charts (Left & Right Sensor)', level=2)
add_side_by_side(
    'Two Chart.js line charts are set up — one cyan (left sensor) and one red (right sensor). '
    'Both start pre-filled with 100 zeroes.\n\n'
    '• animation: false — Disabling animation makes chart updates instant without any lag.\n'
    '• pointRadius: 0 — Hides individual data dots so it looks like a smooth waveform.\n'
    '• tension: 0.4 — Makes the line curve smoothly instead of being jagged/sharp.',
    'const liveChartLeft = new Chart(ctxLiveLeft, {\n    type: \'line\',\n    data: {\n        labels: Array.from({length: 100}, (_, i) => i),\n        datasets: [{\n            data: Array(100).fill(0),\n            borderColor: \'#00f0ff\',\n            borderWidth: 2,\n            pointRadius: 0,\n            tension: 0.4\n        }]\n    },\n    options: {\n        animation: false,\n        // ...\n    }\n});',
    file_ref='index.html', lines_ref='388 – 452'
)

# 3.6 Fleet table
add_heading('3.6  Fleet Overview Table & Bar Chart', level=2)
add_side_by_side(
    'The updateHistoryUI() function reads historical data and populates the table.\n\n'
    '• Line 713: Calculates the minimum diameter across all 4 wheels (FL, FR, BL, BR) to find the worst wheel.\n'
    '• Line 728: The threshold check — any min diameter at or below 123mm triggers a red WARNING status.\n'
    '• Lines 750–758: Dynamically creates and injects a new table row (<tr>) for each OHT record.\n'
    '• Lines 765–770: Updates the bar chart with the same data, coloring bars cyan (healthy) or red (warning).',
    'function updateHistoryUI() {\n    recentData.forEach(row => {\n        const fl = parseFloat(row.FL_mm);\n        const fr = parseFloat(row.FR_mm);\n        const bl = parseFloat(row.BL_mm);\n        const br = parseFloat(row.BR_mm);\n        // Worst (smallest) wheel\n        const minVal = Math.min(fl, fr, bl, br);\n\n        // Threshold check\n        const isWarning = minVal <= 123.0;\n        const statusText = isWarning\n            ? \'WARNING (REPLACE)\'\n            : \'HEALTHY\';\n    });\n}',
    file_ref='index.html', lines_ref='700 – 771'
)

# 3.7 OHT trend
add_heading('3.7  OHT Analysis — Wear Trend Chart', level=2)
add_side_by_side(
    'The updateOhtTrendChart() function filters data for a specific OHT and plots all 4 wheels over time.\n\n'
    '• Line 816: Draws the flat dashed red threshold line at 123mm across the entire chart.\n'
    '• Lines 818–828: Auto-scales the Y-axis to zoom in on the actual wear values, so the trend is always clearly visible.\n'
    '• Lines 848–852: Updates the 4 individual per-wheel charts (FL, FR, BL, BR) with the same filtered data.',
    '// Draw threshold line at 123mm\nohtTrendChart.data.datasets[0].data\n    = ohtData.map(() => 123.0);\n\n// Auto-scale Y axis\nconst minVal = Math.min(...allNumericData, 123.0);\nconst maxVal = Math.max(...allNumericData, 123.0);\nohtTrendChart.options.scales.y.min\n    = Math.floor(minVal - padding);\nohtTrendChart.options.scales.y.max\n    = Math.ceil(maxVal + padding);\n\nohtTrendChart.update();',
    file_ref='index.html', lines_ref='773 – 853'
)

# 3.8 Live polling
add_heading('3.8  The Live Polling Interval', level=2)
add_side_by_side(
    'This single line is what makes the entire dashboard feel "live".\n\n'
    'setInterval() tells the browser to call fetchLive() every 100 milliseconds (10 times per second). '
    'Each time, it fetches the latest sensor readings from /api/live and updates the charts — '
    'all without ever refreshing the page.\n\n'
    'Think of it like your browser sending a text to the Python backend 10 times a second asking '
    '"got any new numbers?" and the backend replying instantly.',
    '// Poll live data 10 times a second\nsetInterval(fetchLive, 100);\n\n// fetchLive sends a request to\n// the backend and updates charts:\nasync function fetchLive() {\n    const res  = await fetch(\'/api/live\');\n    const data = await res.json();\n\n    liveChartLeft.data.datasets[0].data\n        = data.left.history.map(\n            d => d === 999.0 ? 0 : 125.0 + d\n          );\n    liveChartLeft.update();\n}',
    file_ref='index.html', lines_ref='651 – 670 (function) & Line 856 (trigger)'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4: FULL SYSTEM FLOW
# ══════════════════════════════════════════════════════════════════════════════
add_heading('4. Complete System Flow Summary', level=1, color='005B96')
add_para(
    'The diagram below shows how data travels from the physical sensor all the way to the browser dashboard:',
    size=10
)

flow_tbl = doc.add_table(rows=1, cols=1)
flow_tbl.style = 'Table Grid'
flow_cell = flow_tbl.cell(0, 0)
shade_cell(flow_cell, '1E1E2E')
flow_run = flow_cell.paragraphs[0].add_run(
    'OMRON ZP-EIP Sensor  (192.168.250.1:64000)\n'
    '  │\n'
    '  │  app.py Lines 39–61:   Send "MS,01,0" → decode hex → get mm value\n'
    '  ▼\n'
    'app.py Lines 82–119:  Wheel detection (threshold ±40mm)\n'
    '  │\n'
    '  ├─→  app.py Lines 113–115:  Save row to OHT_Wheel_Predictive_Maintenance_Log.csv\n'
    '  │\n'
    '  └─→  app.py Lines 28–33:   Update distance_history (last 100 readings)\n'
    '              │\n'
    '              ▼\n'
    '  app.py Lines 129–155:  Flask API → /api/live  &  /api/history\n'
    '              │\n'
    '              ▼\n'
    '  index.html Line 856:   Browser polls /api/live every 100ms\n'
    '  index.html Lines 651–698:  fetchLive() / fetchHistory() parse data\n'
    '  index.html Lines 388–853:  Chart.js renders live waveforms & trend charts\n'
    '              │\n'
    '              ▼\n'
    '  Dashboard: Live graph  |  Fleet table  |  OHT Wear Trend'
)
flow_run.font.name  = 'Courier New'
flow_run.font.size  = Pt(9)
flow_run.font.color.rgb = RGBColor(0xCD, 0xEF, 0xFF)

doc.add_paragraph()
add_para(
    'Bottom line: This is a full end-to-end Predictive Maintenance monitoring system — '
    'from a physical laser sensor on a factory floor, all the way to a live browser dashboard — '
    'built entirely in Python and standard web technologies.',
    italic=True, size=10
)

# ── Save ───────────────────────────────────────────────────────────────────────
output_path = r'C:\Users\yiron\Desktop\FYProject\Report_Builder\drafts\ProjectODIN_Technical_Breakdown.docx'
doc.save(output_path)
print(f"Saved: {output_path}")
