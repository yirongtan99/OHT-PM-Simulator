from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# â”€â”€ Page Setup â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# â”€â”€ Styles â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
normal_style = doc.styles['Normal']
normal_style.font.name = 'Calibri'
normal_style.font.size = Pt(11)
normal_style.paragraph_format.space_after = Pt(6)
normal_style.paragraph_format.line_spacing = Pt(14)

h1_style = doc.styles['Heading 1']
h1_style.font.name = 'Calibri'
h1_style.font.size = Pt(14)
h1_style.font.bold = True
h1_style.font.color.rgb = RGBColor(0x00, 0x5B, 0x96)
h1_style.paragraph_format.space_before = Pt(18)
h1_style.paragraph_format.space_after  = Pt(6)

h2_style = doc.styles['Heading 2']
h2_style.font.name = 'Calibri'
h2_style.font.size = Pt(12)
h2_style.font.bold = True
h2_style.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
h2_style.paragraph_format.space_before = Pt(12)
h2_style.paragraph_format.space_after  = Pt(4)

# â”€â”€ Helpers â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C0C0C0')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

def add_body(text, indent=False):
    p = doc.add_paragraph(text)
    p.style = 'Normal'
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(0.75 + level * 0.5)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p

def add_formula_box(text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, 'EBF3FB')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name  = 'Courier New'
    run.font.size  = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    doc.add_paragraph()

def add_data_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, '1F497D')
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        bg = 'F5F8FF' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            shade_cell(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            # Bold grand total rows
            if 'GRAND TOTAL' in str(val) or 'Grand Total' in str(val):
                run.bold = True

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# COVER PAGE
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
doc.add_paragraph()
doc.add_paragraph()

cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = cover_title.add_run('FINAL INTERNSHIP REPORT')
run.font.name  = 'Calibri'
run.font.size  = Pt(24)
run.font.bold  = True
run.font.color.rgb = RGBColor(0x00, 0x5B, 0x96)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Professional Attachment â€” AY2025/26')
run.font.name  = 'Calibri'
run.font.size  = Pt(14)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph()
add_horizontal_rule(doc)
doc.add_paragraph()

# Info block
info_table = doc.add_table(rows=7, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ('Author',           'Tan Yi Rong'),
    ('Matriculation No.','U2320967K'),
    ('Company',          'GlobalFoundries Singapore'),
    ('Department',       'F7 Automated Material Handling System (AMHS)'),
    ('Role',             'Equipment Engineer Intern'),
    ('Period',           'May 25, 2026 â€“ Aug 7, 2026 (11 Weeks)'),
    ('Date',             '[Insert Submission Date]'),
]
for r_idx, (label, value) in enumerate(info_data):
    lc = info_table.rows[r_idx].cells[0]
    vc = info_table.rows[r_idx].cells[1]
    lc.width = Cm(5)
    vc.width = Cm(10)
    shade_cell(lc, 'EBF3FB')
    lr = lc.paragraphs[0].add_run(label)
    lr.bold = True; lr.font.size = Pt(11); lr.font.name = 'Calibri'
    lr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    vr = vc.paragraphs[0].add_run(value)
    vr.font.size = Pt(11); vr.font.name = 'Calibri'

doc.add_page_break()

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# ABSTRACT
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
doc.add_heading('Abstract', level=1)
add_horizontal_rule(doc)
add_body(
    'This report documents the work done during an 11-week Professional Attachment at GlobalFoundries '
    'Singapore, within the F7 Automated Material Handling System (AMHS) department. Two workstreams ran '
    'concurrently throughout the attachment. The first and primary focus was Project Odin â€” a sensor-based '
    'initiative designed to replace the manual, vernier caliper-based wheel measurement process currently '
    'performed by line support during OHT Preventive Maintenance (PM), with a consistent and automated '
    'displacement sensor system. The second was the early-stage feasibility study for a Final Year Project '
    '(FYP), exploring a phased approach to automating the manual PM workflow for OHT Mark 2 obstacle '
    'detection sensors. This report covers the problem being solved, the sensing approach, physical setup, '
    'calibration, testing, and future implementation plan for Project Odin, followed by a summary of the '
    'vendor research, cost analysis, and the three-phase automation roadmap for the FYP.'
)

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# SECTION 1 â€” INTRODUCTION
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
doc.add_heading('1.  Introduction & Background', level=1)
add_horizontal_rule(doc)
add_body(
    'GlobalFoundries operates an extensive Automated Material Handling System (AMHS) in its F7 fabrication '
    'facility. The Overhead Hoist Transport (OHT) fleet moves wafers across the cleanroom automatically, '
    'running continuously throughout production. Keeping these vehicles in good working order requires regular '
    'Preventive Maintenance carried out by the line support team â€” the technicians responsible for day-to-day '
    'maintenance and servicing of the OHT fleet.'
)
add_body(
    'Two related problems were identified going into this attachment. The first was with how OHT wheel '
    'measurements were being captured during PM. The second was with the manual, judgment-dependent process '
    'used to verify the OHT obstacle detection sensors. Both share a common root cause â€” over-reliance on '
    'individual technician execution, which introduces variability and inconsistency into what should be a '
    'standardised process.'
)
add_body(
    'Project Odin addressed the first problem. The FYP feasibility study focused on the second, laying out '
    'a phased automation roadmap to reduce manual steps in the sensor PM workflow.'
)

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# SECTION 2 â€” PROJECT ODIN
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
doc.add_heading('2.  Project Odin â€” OHT Wheel Wear Monitoring', level=1)
add_horizontal_rule(doc)

doc.add_heading('2.1  The Problem with the Current Method', level=2)
add_body(
    'As part of the existing PM process, line support is required to measure the diameter of each OHT wheel '
    'using a vernier caliper. While this is a standard measurement tool, applying it consistently in a busy '
    'maintenance environment is harder than it sounds.'
)
add_body(
    'The measurement outcome varies depending on how the caliper is positioned, how much pressure is applied, '
    'and how experienced the person doing it is. Different line support staff measure slightly differently, '
    'resulting in data that is inconsistent between technicians. A more telling sign of the problem is that '
    'the recorded data sometimes shows wheel diameters increasing over time â€” which physically does not make '
    'sense, since a wheel can only wear down, not grow. This pattern in the data is a clear indicator that '
    'the measurements are being affected by human error rather than reflecting the actual condition of the wheels.'
)
add_body(
    'The consequence is that the data collected through the current method cannot be relied upon for any '
    'meaningful trend analysis. Decisions about wheel replacement are still largely based on schedule or visual '
    'inspection rather than actual measured wear.'
)

doc.add_heading('2.2  Objective', level=2)
add_body(
    'Project Odin was brought in to solve this directly. The goal is to replace the vernier caliper '
    'measurement step with a sensor-based system that captures wheel diameter automatically and consistently '
    'every time an OHT passes through the measurement station â€” removing the human variable from the equation '
    'entirely.'
)
add_body(
    'With a reliable, repeatable dataset in place, the department will eventually be able to move from '
    'scheduled wheel replacement to a predictive maintenance model â€” where replacement decisions are driven '
    'by actual measured wear data rather than fixed time intervals.'
)

doc.add_heading('2.3  How the Sensing Works', level=2)
add_body(
    'The system uses a displacement sensor mounted in a fixed position above the OHT track, facing downward. '
    'As an OHT rolls beneath it, the sensor continuously scans the surface below, capturing distance readings '
    'in real time as the wheel passes through its field of view.'
)
add_body(
    'What the sensor captures across each pass is a profile of the wheel\'s curvature as it moves underneath. '
    'The reading starts shallow at the leading edge of the wheel, climbs as the highest point of the wheel '
    'comes into range, then drops back down as the wheel exits. This produces a continuous curve of distance '
    'measurements for each pass.'
)
add_body(
    'From this curve, the system selects the highest valued data point â€” the peak of the curvature, which '
    'corresponds to the closest the wheel surface gets to the sensor. This is the most representative value '
    'of the wheel\'s current diameter. The final calculation is:'
)
add_formula_box('Wheel Diameter  =  125 mm  âˆ’  Peak Sensor Reading')
add_body(
    'The sensor is calibrated against a reference wheel with a known diameter of 125mm. At this reference '
    'the peak reading is 0mm. As a wheel wears and its diameter shrinks, it sits slightly lower on the track, '
    'causing the peak reading to increase â€” and the calculated diameter to drop below 125mm. Wheels measuring '
    'at or below 123mm are flagged for replacement, representing a 2mm reduction from the healthy baseline.'
)

doc.add_heading('2.4  Physical Setup', level=2)
add_body('The installation on the track consists of the following components:')
add_bullet('Displacement Sensor (LiDAR) â€” Mounted above the track, facing downward. Scans continuously as each OHT rolls beneath it.')
add_bullet('Master Amplifier Controller â€” Processes the primary sensor signal.')
add_bullet('Slave Amplifier Controller â€” Mirrors the data for the secondary channel, allowing future dual-sensor coverage.')
add_bullet('Ethernet Connectivity â€” Links the controllers to the data collection system.')
add_bullet('Power Supply Unit â€” Provides stable power to the sensor and controllers.')
add_bullet('Circuit Breaker â€” For safe isolation of the setup.')
doc.add_paragraph()
add_body('The layout is compact and sits unobtrusively on the track without disrupting normal OHT operations.')

doc.add_heading('2.5  Calibration & Testing', level=2)
add_body(
    'Before data collection could start, the sensor was calibrated against a reference wheel with a confirmed '
    'diameter of 125mm. This set the zero point for the system, making all subsequent readings relative to '
    'that known baseline.'
)
add_body(
    'Testing was done by physically rolling a wheel beneath the sensor to simulate an OHT pass. The main '
    'things being verified were that the system could correctly detect the start and end of a wheel pass, '
    'capture the full profile curvature, and reliably select the peak value.'
)
add_body(
    'Several rounds of threshold adjustment were needed before the detection was consistently stable. The '
    'challenge was ensuring the sensor only triggered on actual wheels and not on background objects or noise '
    'at track level. Once the thresholds were confirmed, the system produced reliable results across repeated '
    'test runs, and the calculated diameter values matched manual reference measurements taken alongside.'
)

doc.add_heading('2.6  Data Dashboard & Visualisation', level=2)
add_body(
    'Alongside the sensor hardware, a local web-based dashboard was developed to visualise the collected '
    'data in real time. The dashboard runs as a local web application accessible via any browser on the '
    'same network, giving line support and engineers an easy way to monitor wheel conditions without '
    'opening raw data files. It is organised into four tabs, each serving a different purpose.'
)

# Dashboard tab descriptions as styled boxes
tabs = [
    ('Live Monitoring',
     'Displays a real-time scrolling graph of the sensor\'s current distance readings as each OHT rolls '
     'past the station. The latest OHT number detected and the current sensor reading are shown at the '
     'top of the page. This tab is primarily used during testing and active monitoring sessions to '
     'confirm the sensor is picking up readings correctly.'),
    ('Fleet Overview',
     'Loads all historical wheel measurement data and presents it in a bar chart and a sortable data '
     'table. Each row shows the OHT vehicle number, the minimum recorded wheel diameter across all '
     'four wheels, the timestamp, and a health status â€” HEALTHY or WARNING (REPLACE). Vehicles with '
     'any wheel at or below the 123mm threshold are automatically highlighted in red. The table is '
     'sortable by any column, making it easy to scan the fleet for vehicles that need attention.'),
    ('OHT Analysis',
     'Allows the user to select a specific OHT vehicle and view the wear trend of all four wheels '
     'plotted over time. A dashed red threshold line at 123mm runs across the chart as a visual '
     'reference. If a vehicle is flagged in the Fleet Overview, the engineer can use this tab to '
     'see exactly how wear has been progressing across each individual wheel position '
     '(FL, FR, BL, BR).'),
    ('Side-by-Side Wheel Comparison',
     'Provides an individual chart for each of the four wheel positions, allowing a more granular '
     'breakdown of each wheel\'s degradation state. Charts can be toggled on or off independently '
     'and the layout adjusts automatically. This view is useful when comparing specific wheels to '
     'assess relative wear rates between positions.'),
]
for tab_title, tab_desc in tabs:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    shade_cell(cell, 'F0F6FF')
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(tab_title + '\n')
    r1.bold = True; r1.font.size = Pt(11); r1.font.name = 'Calibri'
    r1.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    r2 = p1.add_run(tab_desc)
    r2.font.size = Pt(11); r2.font.name = 'Calibri'
    doc.add_paragraph()

add_body(
    'All data displayed in the dashboard is pulled directly from the same CSV log file that the sensor '
    'system writes to, keeping the data source consistent across all views.'
)

doc.add_heading('2.7  Data Collected', level=2)
add_body('Each time an OHT passes the sensor station, the system logs one record containing:')
add_bullet('Timestamp of the wheel pass')
add_bullet('OHT vehicle number')
add_bullet('Measured diameter for each of the four wheel positions â€” Front Left (FL), Front Right (FR), Back Left (BL), and Back Right (BR)')
doc.add_paragraph()
add_body(
    'In the current prototype stage, only one physical sensor is deployed. The four-wheel data is generated '
    'using the single sensor\'s reading as the primary value, with minor simulated variation added for the '
    'other three positions. This is a temporary measure to allow the data structure and logging framework to '
    'be fully validated before full deployment.'
)

doc.add_heading('2.8  Future Implementation', level=2)
add_body(
    'The immediate next step is to move the sensor system into the maintenance room inside the fab â€” the '
    'controlled space where OHTs are brought in for scheduled PM. This is the most practical deployment '
    'location because every OHT that comes in for PM will naturally pass through the measurement station, '
    'ensuring regular readings for every vehicle without any additional process burden on the line support team.'
)
add_body(
    'From there, the focus shifts to building up a larger dataset over time. A single measurement every five '
    'to six months per vehicle is not enough to draw reliable wear trend conclusions. The dataset needs to '
    'grow across multiple PM cycles for a significant portion of the fleet before patterns become clear â€” '
    'such as which vehicles wear faster, which wheel positions degrade first, or whether specific track '
    'sections accelerate wear.'
)
add_body(
    'The key challenge in doing this accurately at scale is vehicle identification. Currently, tagging each '
    'measurement to the correct OHT number still requires a manual input step. To make this seamless and '
    'error-free, the planned next development for Project Odin is to integrate an RFID tagging component '
    'into the system. Each OHT would carry an RFID tag, and a reader positioned alongside the sensor would '
    'automatically identify the vehicle as it passes â€” ensuring every measurement is correctly attributed '
    'without any manual input.'
)
add_body(
    'Once RFID tagging is in place and the dataset is large enough, the system will be positioned to support '
    'genuine predictive maintenance â€” flagging specific vehicles for wheel replacement based on their actual '
    'measured wear trajectory rather than a fixed service schedule.'
)

doc.add_page_break()

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# SECTION 3 â€” FYP
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
doc.add_heading('3.  FYP Research & Feasibility â€” Automating the OHT Sensor PM', level=1)
add_horizontal_rule(doc)

doc.add_heading('3.1  Background & Problem Statement', level=2)
add_body(
    'The PM process for the OHT Mark 2 obstacle detection sensors is one of the more involved maintenance '
    'tasks that the line support team handles. This PM is critical from a safety standpoint â€” the sensors '
    'being tested are responsible for detecting obstacles in the OHT\'s path and triggering the vehicle to '
    'stop before any collision occurs. Getting the PM right matters.'
)
add_body('The sensors covered in each PM cycle are the Vehicle Detection Sensor (VHL) and three obstacle sensors â€” OBS Left, OBS Right, and OBS Center.')
add_body('The current process works as follows:')
add_bullet('Line support connects a data cable to the first sensor port and opens the Hokuyo application to view the live sensor waveform.')
add_bullet('Based on what they see on screen, they assess whether the sensor reading is within the acceptable operating range.')
add_bullet('If within spec, they take a screenshot and document it. If not, they manually adjust the physical sensor position and recheck.')
add_bullet('If the sensor cannot be brought into spec after adjustment, it is swapped for a replacement unit and the full verification is repeated.')
add_bullet('This is done for each of the four sensors in sequence â€” meaning the data cable is physically unplugged from one port and replugged into the next, repeatedly.')
add_bullet('Line support also manually shifts a panel or reflective plate in front of each sensor to simulate an obstacle, confirms detection, then removes it to confirm clearance.')
doc.add_paragraph()
add_body(
    'Each full PM cycle takes around one hour per OHT, and the team handles about four OHTs per day. The '
    'process is effective but relies heavily on individual technician judgment â€” both in reading the waveform '
    'and in positioning the panel. These are the gaps the FYP aims to close.'
)

doc.add_heading('3.2  Phased Automation Roadmap', level=2)
add_body(
    'Rather than automating everything at once, the FYP adopts a phased approach that addresses the most '
    'impactful manual steps first and builds towards a fully integrated solution over time.'
)

# Phase boxes
phases = [
    ('Phase 1 â€” Automate the Cable Plug-In/Out',
     'Use a Collaborative Robot (cobot) arm to handle the physical plugging and unplugging of the data cable '
     'between sensor ports. This removes the most repetitive mechanical action in the PM workflow. Automating '
     'this step is estimated to save roughly 10 minutes per PM cycle â€” adding up to approximately 416 '
     'technician hours saved per year across the fleet (30 min x 4 OHTs/day x 208 working days).'),
    ('Phase 2 â€” Automate the Panel Shifting',
     'Introduce a motorised or actuator-driven mechanism to handle the manual panel shifting currently done '
     'to simulate obstacle detection. The mechanism needs to reliably position the panel at the correct '
     'location and angle for each sensor test, then retract it again. This phase has more mechanical '
     'complexity and builds on Phase 1 being stable first.'),
    ('Phase 3 â€” Integrated Stationary Workstation',
     'Bring together the cobot arm, panel shifting mechanism, and waveform verification into a single '
     'integrated workstation. Line support would bring an OHT to the station and step through the PM '
     'sequence by pressing a button at each stage. The system handles the physical actions; the technician '
     'supervises and confirms. This creates a consistent, repeatable workflow with standardised documentation '
     'generated automatically at the end of each session.'),
]
for title, desc in phases:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    shade_cell(cell, 'EBF3FB')
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(title + '\n')
    r1.bold = True; r1.font.size = Pt(11); r1.font.name = 'Calibri'
    r1.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    r2 = p1.add_run(desc)
    r2.font.size = Pt(11); r2.font.name = 'Calibri'
    doc.add_paragraph()

doc.add_heading('3.3  Vendor Research & Cost Findings', level=2)
add_body('Four cobot options were evaluated for Phase 1:')
add_data_table(
    headers=['Vendor', 'Est. Price', 'Reach', 'Vision'],
    rows=[
        ['OMRON TM5S',  'SGD 100k â€“ 150k', '900mm',         'Built-in'],
        ['UR7e',        '~SGD 29k â€“ 36k',  '850mm',         '3rd Party'],
        ['ABB GoFa',    '~SGD 20k â€“ 30k',  'Up to 1300mm',  'SICK (~SGD 5k)'],
        ['ABB PoWa',    '~SGD 15k â€“ 20k',  'Up to 1300mm',  '3rd Party'],
    ],
    col_widths=[4.5, 4, 3.5, 4]
)
add_body(
    'The OMRON was ruled out based on cost â€” at over SGD 100k, the price is too high relative to what Phase 1 '
    'alone delivers. The ABB series came out as the most viable option. One practical challenge identified was '
    'that the existing 4-port plug arrangement on the OHT is not straightforward for a basic vision system to '
    'align to reliably. As a more cost-effective alternative, a proposal was put forward to retrofit all 304 '
    'OHTs with a single centralised magnetic connector. Because a magnetic plug snaps into place by physical '
    'feel, it removes the need for precise vision-guided alignment. At the workstation side, a custom PCB '
    'relay board routes the signal to each sensor in sequence, reproducing the same test workflow from one '
    'fixed connection point.'
)
add_body('Estimated retrofit cost for the full fleet of 304 vehicles:')
add_data_table(
    headers=['Component', 'Unit Cost', 'Total (304 OHTs)'],
    rows=[
        ['Rosenberger 15-pin Magnetic Plug (Premium)',  '~SGD 70',  '~SGD 21,280'],
        ['Standard Pogo-Pin Plug (Budget)',              '~SGD 4',   '~SGD 1,216'],
        ['Custom PCB Bridge Board',                      '~SGD 3',   '~SGD 912'],
        ['Grand Total â€” Premium Route',                  'â€”',        '~SGD 22,192'],
        ['Grand Total â€” Budget Route',                   'â€”',        '~SGD 2,128'],
    ],
    col_widths=[8, 3.5, 4]
)
add_body(
    'One-time retrofit labour at ~SGD 24.34/hr (AE rate based on USD 45k/yr salary, 4-day 12-hour swing '
    'shift): ~SGD 7,344 for 304 OHTs at 1 hour each. Estimated rollout: 2.5 to 4 months at 4â€“6 OHTs per day.'
)

doc.add_heading('3.4  Recommended Path & ROI', level=2)
add_body(
    'The recommended approach is an ABB cobot (~SGD 20k) combined with the budget magnetic plug retrofit '
    '(~SGD 2,128) and a custom PCB relay board â€” bringing the total estimated investment to around SGD 30k. '
    'This is roughly one-fifth the cost of the OMRON option for the same functional Phase 1 outcome.'
)
add_body(
    'From an ROI standpoint, the system is projected to recover approximately SGD 10,100 in direct manpower '
    'costs per year (416 hrs x SGD 24.34/hr), giving a payback period of around three years. Beyond the '
    'direct savings, the automation removes the risk of port damage from repeated manual connections, '
    'eliminates waveform reading inconsistency between operators, and lays the groundwork for Phases 2 '
    'and 3 of the full PM automation roadmap.'
)

doc.add_page_break()

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# SECTION 4 â€” REFLECTIONS
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
doc.add_heading('4.  Reflections', level=1)
add_horizontal_rule(doc)
add_body(
    'Working across both projects during this attachment gave a grounded view of what engineering problem-solving '
    'actually looks like when working within real operational constraints.'
)
add_body(
    'For Project Odin, the most interesting part was understanding why the existing measurement method was '
    'failing â€” not because the vernier caliper is a bad tool, but because the conditions it was being used in '
    'made consistent results nearly impossible. Once that was clear, the case for an automated sensor-based '
    'approach became straightforward. The practical work of setting up, calibrating, and testing the system '
    'was a good exercise in patience â€” getting the detection thresholds right took more iterations than '
    'expected, and there were moments where what looked like a working result turned out to have edge cases '
    'that needed fixing.'
)
add_body(
    'For the FYP, going through the phased planning process was the most useful part. It forced a more '
    'realistic look at what could actually be done within budget and within a reasonable timeframe, rather '
    'than trying to design the ideal solution from scratch. The magnetic plug idea was a direct result of '
    'that â€” finding a way to make the hardware problem simpler instead of adding more expensive vision '
    'hardware to compensate for a difficult-to-align connector.'
)

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# SECTION 5 â€” CONCLUSION
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
doc.add_heading('5.  Conclusion & Next Steps', level=1)
add_horizontal_rule(doc)
add_body(
    'Project Odin now has a validated prototype in place and is actively collecting wheel measurement data. '
    'The next steps are to deploy the system in the maintenance room for proper fleet-wide coverage, and to '
    'integrate RFID tagging so that vehicle identification is fully automatic. Building up the dataset over '
    'time remains the priority before any predictive maintenance decisions can be made with confidence.'
)
add_body(
    'For the FYP, the feasibility study confirms that automating the sensor PM process is viable at a '
    'justifiable cost. Phase 1 â€” the cobot arm for cable plug-in/out â€” is the immediate focus, with the '
    'magnetic plug retrofit as the enabling hardware change. From there, Phase 2 (panel shifting automation) '
    'and Phase 3 (integrated workstation) provide a clear path towards a fully streamlined, '
    'technician-supervised PM workflow that is consistent, documented, and scalable across the fleet.'
)

# â”€â”€ Save â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
output = r'C:\Users\yiron\Desktop\FYProject\Report_Builder\drafts\Final_Report_v3.docx'
doc.save(output)
print(f"Saved: {output}")

